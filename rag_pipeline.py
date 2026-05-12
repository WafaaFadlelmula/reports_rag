#!/usr/bin/env python3
"""
RAG pipeline — ask questions across PDF and Word reports.

One-time indexing: extracts text from all PDFs and .docx files in ./reports, chunks it into
~500-token segments (50-token overlap), embeds with all-MiniLM-L6-v2, and
stores in a local ChromaDB database.

Chat loop: each question retrieves the top-5 most relevant chunks, which are
sent along with the question to GPT for an answer.
"""

import os
import re
import sys
from datetime import datetime
from pathlib import Path

import fitz  # PyMuPDF
import docx   # python-docx
import chromadb
from sentence_transformers import SentenceTransformer
import anthropic

# ── Configuration ──────────────────────────────────────────────────────────────
REPORTS_DIR   = Path("./reports")
CHROMA_DIR    = Path("./chroma_db")
COLLECTION    = "pdf_reports"
CHUNK_TOKENS  = 500   # approximate target chunk size in tokens
OVERLAP_TOKENS = 50   # approximate overlap between consecutive chunks
TOP_K         = 15    # chunks to retrieve per question
CLAUDE_MODEL  = "claude-opus-4-6"
EMBED_MODEL   = "all-MiniLM-L6-v2"
EMBED_BATCH   = 64    # sentences per embedding batch
CHROMA_BATCH  = 512   # documents per ChromaDB upsert batch
BOOST_K       = 3     # extra chunks guaranteed per explicitly named project

# Maps question keywords → a function that recognises matching source filenames.
# When a keyword appears in the question, BOOST_K chunks from matching sources
# are added on top of the normal TOP_K semantic results.
PROJECT_KEYWORDS: dict[str, callable] = {
    "ecotap":  lambda s: "ECOTAP"  in s.upper() or "ECCOTAP" in s.upper(),
    "ecoran":  lambda s: "ECORAN"  in s.upper(),
    "ecoice":  lambda s: "ECOICE"  in s.upper(),
    "bt":      lambda s: "C-PON" in s.upper() or ("AR" in s.upper() and "VR" in s.upper()),
    "testbed": lambda s: "C-PON" in s.upper() or ("AR" in s.upper() and "VR" in s.upper()),
}

# Words-per-token ratio used to convert token targets into word counts.
# all-MiniLM-L6-v2 tokenises roughly 1 word ≈ 1.25 tokens on English prose,
# so 1 token ≈ 0.75 words.
WORDS_PER_TOKEN = 0.75


# ── Text utilities ─────────────────────────────────────────────────────────────

def chunk_text(text: str, source: str) -> list[dict]:
    """
    Split *text* into overlapping word-based chunks that approximate
    CHUNK_TOKENS / OVERLAP_TOKENS token sizes.

    Returns a list of dicts with keys: id, text, source.'
    """
    words         = text.split()
    chunk_words   = max(1, int(CHUNK_TOKENS   * WORDS_PER_TOKEN))
    overlap_words = max(0, int(OVERLAP_TOKENS * WORDS_PER_TOKEN))

    chunks = []
    start  = 0
    idx    = 0

    while start < len(words):
        end        = min(start + chunk_words, len(words))
        chunk_body = " ".join(words[start:end])
        chunks.append({
            "id":     f"{source}::chunk_{idx}",
            "text":   chunk_body,
            "source": source,
        })
        idx += 1
        if end >= len(words):
            break
        start = end - overlap_words

    return chunks


# ── Document extraction ────────────────────────────────────────────────────────

def extract_pdf_text(pdf_path: Path) -> str:
    """Return all text from a PDF as a single string."""
    doc = fitz.open(str(pdf_path))
    pages = [page.get_text() for page in doc]
    doc.close()
    return "\n".join(pages)


def extract_docx_text(docx_path: Path) -> str:
    """Return all text from a .docx file as a single string."""
    document = docx.Document(str(docx_path))
    return "\n".join(para.text for para in document.paragraphs if para.text.strip())


def extract_text(file_path: Path) -> str:
    """Dispatch to the correct extractor based on file extension."""
    if file_path.suffix.lower() == ".pdf":
        return extract_pdf_text(file_path)
    return extract_docx_text(file_path)


# ── Indexing ───────────────────────────────────────────────────────────────────

def build_index() -> None:
    """
    Extract text from every PDF and .docx in REPORTS_DIR, chunk it, embed
    the chunks with SentenceTransformer, and persist them in ChromaDB.
    """
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    doc_files = sorted(
        list(REPORTS_DIR.glob("*.pdf")) + list(REPORTS_DIR.glob("*.docx"))
    )

    if not doc_files:
        print(f"[error] No PDF or Word files found in {REPORTS_DIR.resolve()}")
        print("        Add .pdf or .docx files to that folder and re-run.")
        sys.exit(1)

    print(f"Found {len(doc_files)} file(s).  Loading embedding model…")
    embedder = SentenceTransformer(EMBED_MODEL)

    # Prepare ChromaDB (cosine similarity)
    chroma     = chromadb.PersistentClient(path=str(CHROMA_DIR))
    collection = chroma.create_collection(
        name=COLLECTION,
        metadata={"hnsw:space": "cosine"},
    )

    all_ids, all_docs, all_metas = [], [], []
    all_embeds = []

    for pdf_path in doc_files:
        print(f"  • {pdf_path.name}", end="", flush=True)
        raw_text = extract_text(pdf_path)
        chunks   = chunk_text(raw_text, pdf_path.name)
        texts    = [c["text"] for c in chunks]

        # Embed in mini-batches to avoid OOM on large files
        embeds = embedder.encode(
            texts,
            batch_size=EMBED_BATCH,
            show_progress_bar=False,
            normalize_embeddings=True,
        ).tolist()

        print(f"  ({len(chunks)} chunks)")
        for chunk, emb in zip(chunks, embeds):
            all_ids.append(chunk["id"])
            all_docs.append(chunk["text"])
            all_metas.append({"source": chunk["source"]})
            all_embeds.append(emb)

    # Upsert in batches
    total = len(all_ids)
    for i in range(0, total, CHROMA_BATCH):
        collection.add(
            ids        = all_ids[i : i + CHROMA_BATCH],
            documents  = all_docs[i : i + CHROMA_BATCH],
            metadatas  = all_metas[i : i + CHROMA_BATCH],
            embeddings = all_embeds[i : i + CHROMA_BATCH],
        )

    print(f"\nIndexed {total} chunks from {len(doc_files)} file(s) "
          f"→ {CHROMA_DIR.resolve()}\n")


# ── Retrieval ──────────────────────────────────────────────────────────────────

def retrieve(
    question:    str,
    embedder:    SentenceTransformer,
    collection,
    all_sources: list[str],
) -> list[dict]:
    """
    Embed *question* and return the most relevant chunks.

    Two-stage retrieval:
    1. Semantic search — top TOP_K chunks by cosine similarity.
    2. Keyword boost — for every project name found in the question,
       add BOOST_K chunks from that project's sources, even if they
       didn't rank in the top TOP_K.  This prevents named projects
       from being silently omitted when their chunks score slightly
       lower than others.

    Returns a list of dicts with keys: text, source, score.
    """
    q_embed = embedder.encode([question], normalize_embeddings=True).tolist()

    # ── Stage 1: semantic search ───────────────────────────────────────────────
    results = collection.query(
        query_embeddings=q_embed,
        n_results=TOP_K,
        include=["documents", "metadatas", "distances"],
    )

    seen: set[str] = set()
    chunks: list[dict] = []

    def _add(doc, meta, dist):
        key = meta["source"] + doc[:80]
        if key not in seen:
            seen.add(key)
            chunks.append({
                "text":   doc,
                "source": meta["source"],
                "score":  round(1.0 - dist / 2.0, 4),
            })

    for doc, meta, dist in zip(
        results["documents"][0],
        results["metadatas"][0],
        results["distances"][0],
    ):
        _add(doc, meta, dist)

    # ── Stage 2: keyword boost ─────────────────────────────────────────────────
    q_lower = question.lower()
    for keyword, matcher in PROJECT_KEYWORDS.items():
        if keyword not in q_lower:
            continue

        # Which source files match this project keyword?
        matched = [s for s in all_sources if matcher(s)]
        if not matched:
            continue

        # Only boost sources not already fully represented
        already = {c["source"] for c in chunks}
        missing = [s for s in matched if s not in already]
        if not missing:
            continue

        try:
            boosted = collection.query(
                query_embeddings=q_embed,
                n_results=BOOST_K,
                where={"source": {"$in": missing}},
                include=["documents", "metadatas", "distances"],
            )
            for doc, meta, dist in zip(
                boosted["documents"][0],
                boosted["metadatas"][0],
                boosted["distances"][0],
            ):
                _add(doc, meta, dist)
        except Exception:
            pass  # older ChromaDB versions may not support $in filter

    return sorted(chunks, key=lambda x: x["score"], reverse=True)


# ── Generation ─────────────────────────────────────────────────────────────────

def generate_answer(
    question: str,
    chunks:   list[dict],
    client:   anthropic.Anthropic,
) -> str:
    """
    Format retrieved chunks as context, send to Claude, and return the answer.
    """
    context_blocks = []
    for i, c in enumerate(chunks, start=1):
        context_blocks.append(
            f"[Excerpt {i} — {c['source']} (relevance {c['score']})]\n{c['text']}"
        )
    context = "\n\n---\n\n".join(context_blocks)

    system = (
        "You are a professional technical writer and research analyst. "
        "Using ONLY the provided excerpts, write a well-structured, fluent document "
        "that directly addresses the user's request. "
        "Use clear headings and paragraphs. Write in a formal, professional tone. "
        "Cite the source file name in parentheses when drawing on specific information. "
        "If the excerpts do not contain enough detail to cover a point, say so briefly "
        "rather than inventing content."
    )

    user_message = (
        f"EXCERPTS FROM REPORTS:\n\n{context}\n\n"
        f"QUESTION: {question}"
    )

    with client.messages.stream(
        model=CLAUDE_MODEL,
        max_tokens=4096,
        system=system,
        messages=[{"role": "user", "content": user_message}],
    ) as stream:
        return stream.get_final_message().content[0].text


# ── Save helpers ──────────────────────────────────────────────────────────────

OUTPUT_DIR = Path("./outputs")

def _make_filename(question: str) -> str:
    """Derive a short slug from the question plus a timestamp."""
    slug = re.sub(r"[^\w\s-]", "", question.lower())
    slug = re.sub(r"[\s_-]+", "_", slug).strip("_")[:40]
    ts   = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"{slug}_{ts}"


def save_as_md(question: str, answer: str) -> Path:
    """Write the answer to a Markdown file and return the path."""
    OUTPUT_DIR.mkdir(exist_ok=True)
    path = OUTPUT_DIR / (_make_filename(question) + ".md")
    path.write_text(f"# {question}\n\n{answer}\n", encoding="utf-8")
    return path


def save_as_docx(question: str, answer: str) -> Path:
    """Write the answer to a Word document and return the path."""
    OUTPUT_DIR.mkdir(exist_ok=True)
    path = OUTPUT_DIR / (_make_filename(question) + ".docx")

    document = docx.Document()
    document.add_heading(question, level=1)

    # Split on blank lines to preserve paragraph breaks from the model
    for block in answer.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        # Lines starting with ## / # are treated as sub-headings
        if block.startswith("## "):
            document.add_heading(block[3:], level=2)
        elif block.startswith("# "):
            document.add_heading(block[2:], level=2)
        else:
            document.add_paragraph(block)

    document.save(str(path))
    return path


# ── Main chat loop ─────────────────────────────────────────────────────────────

def main() -> None:
    # ── One-time indexing (skip if DB already exists) ──────────────────────────
    if not CHROMA_DIR.exists():
        print("No index found — building it now…\n")
        build_index()
    else:
        print(f"Using existing index at {CHROMA_DIR.resolve()}")

    # ── Load embedder + open DB ────────────────────────────────────────────────
    print("Loading embedding model…")
    embedder   = SentenceTransformer(EMBED_MODEL)
    chroma     = chromadb.PersistentClient(path=str(CHROMA_DIR))
    collection = chroma.get_collection(COLLECTION)
    all_sources = list({m["source"] for m in collection.get(include=["metadatas"])["metadatas"]})

    api_client = anthropic.Anthropic()  # reads ANTHROPIC_API_KEY from env

    doc_count = collection.count()
    print(f"Ready — {doc_count} chunks indexed.")
    print("Type your question, 'save md' / 'save docx' to save the last answer, "
          "or 'quit' to exit.\n")

    # ── Interactive loop ───────────────────────────────────────────────────────
    last_question: str = ""
    last_answer:   str = ""

    while True:
        try:
            user_input = input("You: ").strip()
        except (EOFError, KeyboardInterrupt):
            print("\nGoodbye.")
            break

        if not user_input:
            continue
        if user_input.lower() in {"quit", "exit", "q"}:
            print("Goodbye.")
            break

        # ── Save commands ──────────────────────────────────────────────────────
        if user_input.lower() in {"save md", "save markdown"}:
            if not last_answer:
                print("Nothing to save yet — ask a question first.\n")
            else:
                path = save_as_md(last_question, last_answer)
                print(f"Saved → {path.resolve()}\n")
            continue

        if user_input.lower() in {"save docx", "save word", "save doc"}:
            if not last_answer:
                print("Nothing to save yet — ask a question first.\n")
            else:
                path = save_as_docx(last_question, last_answer)
                print(f"Saved → {path.resolve()}\n")
            continue

        # ── Normal question ────────────────────────────────────────────────────
        chunks = retrieve(user_input, embedder, collection, all_sources)
        if not chunks:
            print("Claude:No relevant content found in the reports.\n")
            continue

        sources = sorted({c["source"] for c in chunks})
        print(f"[Sources: {', '.join(sources)}]")

        print("Claude:", end="", flush=True)
        answer = generate_answer(user_input, chunks, api_client)
        print(answer)
        print()

        last_question = user_input
        last_answer   = answer


if __name__ == "__main__":
    main()
